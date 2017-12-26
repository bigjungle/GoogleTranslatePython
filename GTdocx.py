#!/usr/bin/env python3
#-*- coding:utf-8 -*-
from docx import Document
from selenium import webdriver
import random,time
import os, sys
import threading,inspect

# ---start--- for pyinstaller include chromedriver.exe ,also need edit spec file
current_folder = os.path.realpath(os.path.abspath(os.path.split(inspect.getfile(inspect.currentframe() ))[0]))
chromedriver = os.path.join(current_folder,"chromedriver.exe")
# via this way, you explicitly let Chrome know where to find 
# the webdriver.
b = webdriver.Chrome(executable_path = chromedriver) 
# ---end--- 

#b = webdriver.Chrome();
#b.get("https://translate.google.cn/#auto/en");  # chinese to english
b.get("https://translate.google.cn/?hl=en#en/zh-CN");  #english to chinese

#rootdir = 'E:/phjs/work/source' 
#Change to use relative path 
rootdir = os.path.join(os.getcwd() ,'source' )   #Place the docx files you need to translate here,The program monitors this folder
tardir = os.path.join(os.getcwd() ,'target' )    #The translated file is here 

def translate():
	list = os.listdir(rootdir) #list all folder and files
	for i in range(0,len(list)):
		# path = os.path.join(rootdir,list[i])
		path =  rootdir + '/' + list[i]
		pathtar =  tardir + '/' + list[i]
		if os.path.isfile(path):
			start = time.clock()
			document = Document(path) 
			text = ""
			line = ""
			resulttext = ""
			l = [ paragraph.text.encode('utf-8') for paragraph in document.paragraphs];# use encode to avoid gbk special character error 
			for p in l:
				b.find_element_by_id("source").clear() 
				time.sleep(random.uniform(0.1,0.2))
				str_list=p.decode('utf8').split('\n'); 
				for index in range(len(str_list)):
					if str_list[index] == "":
						print ("empty")
					else:
						b.find_element_by_id("source").clear() 
						b.find_element_by_id("source").send_keys(str_list[index])
						time.sleep(random.uniform(0.1,0.2))
						b.find_element_by_id("gt-submit").click()
						time.sleep(random.uniform(0.1,0.2))
						result = ""
						while result == "" :
							result = b.find_element_by_id("result_box").text
						resulttext = resulttext + result + "\n"
						#print("---------str_list----------");
			print("translate ok")	 
			documentwr = Document()
			p = documentwr.add_paragraph(resulttext)				
			documentwr.save(pathtar)
			print(pathtar + "saved")	
			os.remove(path); #remove source files after translate
			elapsed = (time.clock() - start)
			print("Time used:",elapsed) 
			time.sleep(random.uniform(0.5,1.5)) 
def fun_timer(): 
	translate()
	global timer
	timer = threading.Timer(3, fun_timer) #call translate per 3 seconds
	timer.start()
	
timer = threading.Timer(1, fun_timer)
timer.start()	
# print(result);

#b.quit()
